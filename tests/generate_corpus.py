"""
paper2tex: generate_test_corpus — Create synthetic test .docx files for testing.

Generates minimal .docx files with known content to validate the extraction pipeline.
Each test case targets specific extraction challenges.
"""

from __future__ import annotations

import logging
from pathlib import Path

log = logging.getLogger(__name__)

try:
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def create_test_basic(output_dir: Path) -> Path:
    """Test case 1: Basic paper with sections, paragraphs, and a simple table."""
    doc = Document()
    doc.add_heading("A Study of Basic Paper Structure", level=0)

    # Author info as a paragraph
    p = doc.add_paragraph("John Smith, Jane Doe")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("University of Testing, Test City, TC 12345")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This paper presents a basic study to validate the paper2tex conversion pipeline. "
        "We demonstrate that simple text, tables, and section structure are preserved correctly."
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The problem of converting Word documents to LaTeX has been studied extensively. "
        "In this work, we present a novel approach that uses deterministic extraction."
    )

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Our approach consists of three main steps: extraction, assembly, and verification."
    )

    # Simple table
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Method", "Accuracy", "Time (s)"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Baseline", "85.3%", "12.5"],
        ["Proposed", "92.1%", "8.3"],
        ["Oracle", "99.0%", "1.0"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_paragraph("Table 1: Comparison of methods on the test dataset.")

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "As shown in Table 1, our proposed method achieves 92.1% accuracy, "
        "outperforming the baseline by 6.8 percentage points."
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "We have demonstrated that our approach effectively converts Word documents "
        "to conference-formatted LaTeX files."
    )

    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] A. Smith, B. Jones, \"A survey of document conversion,\" in Proc. ICML, 2024.")
    doc.add_paragraph("[2] C. Lee, D. Wang, \"LaTeX automation for academic papers,\" IEEE Trans., 2023.")
    doc.add_paragraph("[3] E. Brown, \"Deep learning for document processing,\" NeurIPS, 2025.")

    path = output_dir / "test_basic.docx"
    doc.save(str(path))
    log.info("Created: %s", path)
    return path


def create_test_special_chars(output_dir: Path) -> Path:
    """Test case 2: Paper with special characters, units, Greek letters."""
    doc = Document()
    doc.add_heading("Temperature Effects on Material Properties", level=0)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We study materials at temperatures ranging from 25°C to 500°C. "
        "The thermal conductivity coefficient α varies as α ≈ 2.5 × 10⁻³ W/(m·K). "
        "Results show a ±3% margin of error."
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "The relationship between temperature and conductivity follows: "
        "σ = σ₀ × exp(−Eₐ/kT), where σ₀ is the pre-exponential factor, "
        "Eₐ is the activation energy (≈ 0.5 eV), "
        "k is Boltzmann's constant (8.617 × 10⁻⁵ eV/K), "
        "and T is temperature in Kelvin."
    )

    doc.add_heading("Experimental Setup", level=1)
    doc.add_paragraph(
        "Samples were prepared with dimensions 10 µm × 20 µm × 5 µm. "
        "Measurements were performed at ΔT = 50°C intervals. "
        "The resistance was measured in kΩ with precision ≤ 0.1%. "
        "pH values ranged from 6.5–8.0."
    )

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "The Seebeck coefficient S was found to be −42 µV/K ± 2 µV/K. "
        "The figure of merit ZT ≈ 1.2 at 300°C, which is ≥ the theoretical minimum. "
        "The density ρ = 5.32 g/cm³."
    )

    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] M. Thévenot, \"Thermal properties at 500°C,\" J. Appl. Phys., 2024.")

    path = output_dir / "test_special_chars.docx"
    doc.save(str(path))
    log.info("Created: %s", path)
    return path


def create_test_lists_and_code(output_dir: Path) -> Path:
    """Test case 3: Paper with lists, code blocks, and algorithms."""
    doc = Document()
    doc.add_heading("An Algorithm for Document Conversion", level=0)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We present an algorithm for converting Word documents to LaTeX. "
        "Our approach handles bullet lists, numbered lists, and pseudocode."
    )

    doc.add_heading("Approach", level=1)
    doc.add_paragraph("Our pipeline consists of the following steps:")

    # Numbered list
    for i, step in enumerate([
        "Parse the .docx XML structure",
        "Extract math equations deterministically",
        "Reconstruct tables with merged cells",
        "Generate bibliography from references",
        "Compile and verify the output",
    ], 1):
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Key Features", level=2)
    for feature in [
        "Verbatim text preservation — zero rewording",
        "Deterministic math conversion — no LLM involvement",
        "Automatic error correction — up to 10 iterations",
    ]:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("Implementation", level=1)
    doc.add_paragraph("The core extraction function in Python:")

    # Code block (using monospace font)
    code = doc.add_paragraph()
    run = code.add_run(
        "def extract(docx_path: Path) -> dict:\n"
        "    with zipfile.ZipFile(docx_path) as zf:\n"
        "        doc_xml = zf.read('word/document.xml')\n"
        "    tree = etree.fromstring(doc_xml)\n"
        "    return parse_content(tree)\n"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Paper2tex Documentation, 2026.")

    path = output_dir / "test_lists_code.docx"
    doc.save(str(path))
    log.info("Created: %s", path)
    return path


def generate_test_corpus(output_dir: Path) -> list[Path]:
    """Generate all test corpus files.

    Returns list of created file paths.
    """
    if not HAS_DOCX:
        log.error("python-docx is required to generate test corpus: pip install python-docx")
        return []

    output_dir.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []

    try:
        created.append(create_test_basic(output_dir))
    except Exception as e:
        log.error("Failed to create test_basic: %s", e)

    try:
        created.append(create_test_special_chars(output_dir))
    except Exception as e:
        log.error("Failed to create test_special_chars: %s", e)

    try:
        created.append(create_test_lists_and_code(output_dir))
    except Exception as e:
        log.error("Failed to create test_lists_code: %s", e)

    log.info("Generated %d test corpus files in %s", len(created), output_dir)
    return created


def main() -> None:
    """CLI entry point."""
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    output_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("tests/corpus")
    files = generate_test_corpus(output_dir)

    if files:
        print(f"✅ Generated {len(files)} test files:")
        for f in files:
            print(f"   {f}")
    else:
        print("❌ No test files generated. Is python-docx installed?")


if __name__ == "__main__":
    main()
